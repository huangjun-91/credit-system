"""
Word排版工具 - 支持两套格式规范
模式1: 党政机关公文格式 (GB/T 9704-2012)
模式2: 课题研究报告格式（云阳县教育学会规范）
"""

import io
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from flask import Blueprint, render_template, request, send_file, flash

format_bp = Blueprint('format', __name__, url_prefix='/format')

# ============================================================
# 模式1: 公文格式 (GB/T 9704-2012)
# ============================================================
GONGWEN_CONFIG = {
    'name': '公文格式',
    'page_width': 21.0, 'page_height': 29.7,
    'margin_top': 3.7, 'margin_bottom': 3.5,
    'margin_left': 2.7, 'margin_right': 2.7,
    'line_spacing_fixed': 28,
    'title_font': '方正小标宋体', 'title_size': 22, 'title_bold': False,
    'h1_font': '方正黑体体', 'h1_size': 16, 'h1_bold': False,
    'h2_font': '方正楷体体', 'h2_size': 16, 'h2_bold': False,
    'h3_font': '方正仿宋体', 'h3_size': 16, 'h3_bold': False,
    'h4_font': '方正仿宋体', 'h4_size': 16, 'h4_bold': False,
    'h5_font': '方正仿宋体', 'h5_size': 16, 'h5_bold': False, 'h5_prefix': None,
    'body_font': '方正仿宋体', 'body_size': 16,
    'body_first_indent': 2,
    'table_font': '方正仿宋体', 'table_size': 14,  # 表格通常小一号
    'align_body': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'page_number_style': 'official',
}

# ============================================================
# 模式2: 课题报告格式 (云阳县教育学会)
# ============================================================
KETI_CONFIG = {
    'name': '课题报告格式',
    'page_width': 21.0, 'page_height': 29.7,
    'margin_top': 2.54, 'margin_bottom': 2.54,
    'margin_left': 3.18, 'margin_right': 3.18,
    'line_spacing_fixed': 20,          # 固定值20磅
    'para_space_before': 0,
    'para_space_after': 0,

    # 题目：2号或小2号宋体加粗
    'title_font': '宋体', 'title_size': 18, 'title_bold': True,   # 小2=18pt

    # 承担单位：宋体居中（沿用正文设置）

    # 一级标题：黑体三号，序号"一、"
    'h1_font': '黑体', 'h1_size': 16, 'h1_bold': False,
    # 二级标题：黑体小三，序号"（一）"
    'h2_font': '黑体', 'h2_size': 15, 'h2_bold': False,    # 小三=15pt
    # 三级标题：黑体四号，序号"1．"
    'h3_font': '黑体', 'h3_size': 14, 'h3_bold': False,    # 四号=14pt
    # 四级标题：宋体小四，序号"（1）"
    'h4_font': '宋体', 'h4_size': 12, 'h4_bold': False,    # 小四=12pt
    # 五级标题：宋体小四，序号"①"
    'h5_font': '宋体', 'h5_size': 12, 'h5_bold': False, 'h5_prefix': '①',

    # 正文：宋体小四
    'body_font': '宋体', 'body_size': 12,
    'body_first_indent': 2,

    # 表格用字（同正文：宋体小四）
    'table_font': '宋体', 'table_size': 12,       # 小四号

    # 参考文献
    'ref_font': '宋体', 'ref_size': 10.5,         # 五号

    'align_body': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'page_number_style': 'simple',
}

# 层级标题匹配
LEVEL1_RE = re.compile(r'^[一二三四五六七八九十]+[、．](?!\d)')
LEVEL2_RE = re.compile(r'^（[一二三四五六七八九十]+）')
LEVEL3_RE = re.compile(r'^\d+[．\.]')
LEVEL4_RE = re.compile(r'^（\d+）')
LEVEL5_RE = re.compile(r'^①')

# 参考文献检测
REF_RE = re.compile(r'^\[(\d+|[一二三四五六七八九十]+)\]')
# 承担单位检测（常见写法）
DANWEI_RE = re.compile(r'(云阳县|重庆市|重庆).{2,20}(学校|小学|中学|中心校|幼儿园|学院|大学|教育局|委员会)')
# 主要参考文献标题
REF_TITLE_RE = re.compile(r'^[参考文献|参考文献|主要参考文献|参考资料]')


def set_run_font(run, font_name, size_pt, bold=False, color=None):
    """设置run字体 - 完全安全的实现"""
    try:
        run.font.size = Pt(size_pt)
    except Exception:
        pass
    try:
        run.font.bold = bold
    except Exception:
        pass
    try:
        if color:
            run.font.color.rgb = color
    except Exception:
        pass

    # 设置中英文字体（用最安全的方式）
    try:
        rPr = run._element.get_or_add_rPr()
    except Exception:
        try:
            # 如果 get_or_add_rPr 失败，手动构造
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                from lxml import etree
                rPr_elem = OxmlElement('w:rPr')
                run._element.append(rPr_elem)
                rPr = rPr_elem
        except Exception:
            return  # 彻底放弃

    if rPr is None:
        return

    try:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            import lxml.etree as etree
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    except Exception:
        pass


def set_para_fixed_spacing(para, pt_value):
    """设置行距为固定值"""
    try:
        pPr = para._element.get_or_add_pPr()
        if pPr is None:
            return
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), str(int(pt_value * 20)))  # 1pt = 20 twips
        spacing.set(qn('w:lineRule'), 'exact')
    except Exception:
        pass


def format_table(table, config):
    """格式化表格：设置表格内字体、边框"""
    try:
        # 设置表格对齐方式
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 清除段落缩进
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    # 表格内居中
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    for run in para.runs:
                        set_run_font(run, config['table_font'], config['table_size'], False)
                    if not para.runs:
                        run = para.add_run(para.text)
                        set_run_font(run, config['table_font'], config['table_size'], False)

                    # 固定行距
                    set_para_fixed_spacing(para, max(config.get('line_spacing_fixed', 20) - 4, 14))

        # 设置表格边框（0.5磅实线）
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), '4')
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), '000000')
            borders.append(elem)
        # 清除旧边框
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(borders)
    except Exception:
        pass  # 表格格式化失败不中断整体流程


def add_simple_page_number(doc, font_size=10):
    """简单页码：居中阿拉伯数字"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    # 清空
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    set_run_font(run, 'Times New Roman', font_size, False)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._element.append(instrText)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar_end)


class DocumentFormatter:
    def __init__(self, doc, config):
        self.doc = doc
        self.cfg = config

    def set_page_layout(self):
        s = self.doc.sections[0]
        s.page_width = Cm(self.cfg['page_width'])
        s.page_height = Cm(self.cfg['page_height'])
        s.top_margin = Cm(self.cfg['margin_top'])
        s.bottom_margin = Cm(self.cfg['margin_bottom'])
        s.left_margin = Cm(self.cfg['margin_left'])
        s.right_margin = Cm(self.cfg['margin_right'])
        s.footer_distance = Cm(1.0)

    def detect_level(self, text):
        if not text:
            return 0
        if LEVEL1_RE.match(text):
            return 1
        if LEVEL2_RE.match(text):
            return 2
        if LEVEL3_RE.match(text):
            return 3
        if LEVEL4_RE.match(text):
            return 4
        if LEVEL5_RE.match(text):
            return 5
        return 0

    def is_reference_title(self, text):
        """检测是否为参考文献标题行"""
        return bool(REF_TITLE_RE.match(text))

    def is_reference_entry(self, text):
        """检测是否为参考文献条目"""
        return bool(REF_RE.match(text))

    def is_danwei_line(self, text, is_after_title):
        """检测是否为承担单位行"""
        if not is_after_title:
            return False
        if len(text) < 3:
            return False
        # 如果包含"云阳县"、"学校"、"小学"等且内容不长
        if DANWEI_RE.search(text):
            return True
        return False

    def format_paragraphs(self):
        has_met_content = False        # 是否遇到过非空段落
        title_handled = False          # 文件主标题已处理
        danwei_handled = False         # 承担单位已处理
        in_ref_section = False

        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # ---- 参考文献段 ----
            if self.is_reference_title(text):
                in_ref_section = True
            if in_ref_section:
                is_ref_title = self.is_reference_title(text)
                font_name = self.cfg.get('ref_font', self.cfg['body_font'])
                font_size = self.cfg.get('ref_size', self.cfg['body_size'])
                bold = is_ref_title
                alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                first_indent = 0
            else:
                # ---- 检测文本类型 ----
                level = self.detect_level(text)

                # 1) 文件标题：第一个内容的段落（非层级标题、非参考文献）
                is_title = False
                if not has_met_content and level == 0 and not self.is_reference_entry(text):
                    is_title = True
                    has_met_content = True
                    title_handled = True

                # 2) 承担单位：标题后的下一个非空段落（含单位关键词）
                is_danwei = False
                if title_handled and not danwei_handled and level == 0:
                    if self.is_danwei_line(text, True) or (len(text) < 40 and not text.startswith('一')):
                        # 标题后第一个短段落视为承担单位
                        is_danwei = True
                        danwei_handled = True

                # ---- 根据类型设置格式 ----
                if is_title:
                    font_name = self.cfg['title_font']
                    font_size = self.cfg['title_size']
                    bold = self.cfg['title_bold']
                    alignment = WD_ALIGN_PARAGRAPH.CENTER
                    first_indent = 0
                elif is_danwei:
                    font_name = self.cfg['body_font']
                    font_size = self.cfg['body_size']
                    bold = False
                    alignment = WD_ALIGN_PARAGRAPH.CENTER
                    first_indent = 0
                elif level == 1:
                    font_name = self.cfg['h1_font']
                    font_size = self.cfg['h1_size']
                    bold = self.cfg.get('h1_bold', False)
                    alignment = WD_ALIGN_PARAGRAPH.LEFT
                    first_indent = 0
                elif level == 2:
                    font_name = self.cfg['h2_font']
                    font_size = self.cfg['h2_size']
                    bold = self.cfg.get('h2_bold', False)
                    alignment = WD_ALIGN_PARAGRAPH.LEFT
                    first_indent = 0
                elif level == 3:
                    font_name = self.cfg['h3_font']
                    font_size = self.cfg['h3_size']
                    bold = self.cfg.get('h3_bold', False)
                    alignment = WD_ALIGN_PARAGRAPH.LEFT
                    first_indent = 0
                elif level == 4:
                    font_name = self.cfg['h4_font']
                    font_size = self.cfg['h4_size']
                    bold = self.cfg.get('h4_bold', False)
                    alignment = WD_ALIGN_PARAGRAPH.LEFT
                    first_indent = 0
                elif level == 5:
                    font_name = self.cfg['h5_font']
                    font_size = self.cfg['h5_size']
                    bold = self.cfg.get('h5_bold', False)
                    alignment = WD_ALIGN_PARAGRAPH.LEFT
                    first_indent = 0
                else:
                    font_name = self.cfg['body_font']
                    font_size = self.cfg['body_size']
                    bold = False
                    alignment = self.cfg.get('align_body', WD_ALIGN_PARAGRAPH.JUSTIFY)
                    first_indent = self.cfg.get('body_first_indent', 2)

            # ---- 应用格式 ----
            for run in para.runs:
                set_run_font(run, font_name, font_size, bold)
            if not para.runs:
                run = para.add_run(text)
                set_run_font(run, font_name, font_size, bold)

            para.alignment = alignment
            if first_indent > 0:
                para.paragraph_format.first_line_indent = Pt(font_size * first_indent)
            else:
                para.paragraph_format.first_line_indent = Pt(0)

            # 行距
            set_para_fixed_spacing(para, self.cfg['line_spacing_fixed'])
            para.paragraph_format.space_before = Pt(self.cfg.get('para_space_before', 0))
            para.paragraph_format.space_after = Pt(self.cfg.get('para_space_after', 0))

            # 1-3级标题加段前间距
            if not in_ref_section and level in (1, 2, 3):
                para.paragraph_format.space_before = Pt(self.cfg['line_spacing_fixed'])
                para.paragraph_format.space_after = Pt(self.cfg['line_spacing_fixed'])

    def format_tables(self):
        """格式化文档中所有表格"""
        for table in self.doc.tables:
            format_table(table, self.cfg)

    def add_page_number(self):
        style = self.cfg.get('page_number_style', 'simple')
        if style == 'official':
            # 公文页码：一字线 + 页码
            section = self.doc.sections[0]
            footer = section.footer
            footer.is_linked_to_previous = False
            for p in footer.paragraphs:
                p.clear()
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run1 = fp.add_run('—')
            set_run_font(run1, 'Times New Roman', 14, False)
            run2 = fp.add_run()
            set_run_font(run2, 'Times New Roman', 14, False)
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            run2._element.append(fldChar_begin)
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' PAGE '
            run2._element.append(instrText)
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            run2._element.append(fldChar_end)
            run3 = fp.add_run('—')
            set_run_font(run3, 'Times New Roman', 14, False)
        else:
            add_simple_page_number(self.doc, 10)

    def run(self):
        self.set_page_layout()
        self.format_paragraphs()
        self.format_tables()      # 表格排版
        self.add_page_number()
        return self.doc


# ============================================================
# Flask 路由
# ============================================================

CONFIGS = {
    'gongwen': GONGWEN_CONFIG,
    'keti': KETI_CONFIG,
}

@format_bp.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('请选择文件', 'error')
            return render_template('format.html')

        file = request.files['file']
        if file.filename == '':
            flash('请选择文件', 'error')
            return render_template('format.html')
        if not file.filename.lower().endswith('.docx'):
            flash('仅支持 .docx 格式', 'error')
            return render_template('format.html')

        # 选择模式
        mode = request.form.get('mode', 'keti')
        base_config = CONFIGS.get(mode, KETI_CONFIG)
        config = base_config.copy()

        try:
            config['title_size'] = float(request.form.get('title_size', config['title_size']))
            config['body_size'] = float(request.form.get('body_size', config['body_size']))
            config['line_spacing_fixed'] = float(request.form.get('line_spacing', config['line_spacing_fixed']))
            config['margin_top'] = float(request.form.get('margin_top', config['margin_top']))
            config['margin_bottom'] = float(request.form.get('margin_bottom', config['margin_bottom']))
            config['margin_left'] = float(request.form.get('margin_left', config['margin_left']))
            config['margin_right'] = float(request.form.get('margin_right', config['margin_right']))
        except ValueError:
            pass

        try:
            # 验证文件是否为有效的 docx (ZIP格式)
            file_bytes = file.read()
            if len(file_bytes) < 100:
                flash('文件内容为空或文件过小', 'error')
                return render_template('format.html')
            if file_bytes[0:2] != b'PK':
                flash('文件格式错误：不是有效的 .docx 文件。'
                      '请确认文件是用 Word 或 WPS 直接创建的 .docx 格式，'
                      '而不是将 .doc 文件直接改后缀名得到的。', 'error')
                return render_template('format.html')

            # 重新构造文件对象
            file.stream = io.BytesIO(file_bytes)
            file.stream.seek(0)

            doc = Document(file.stream)

            # 打印文档信息协助调试
            para_count = len(doc.paragraphs)
            first_text = doc.paragraphs[0].text.strip() if doc.paragraphs else ''
            print(f"[FORMAT] Processing: {file.filename}, paragraphs={para_count}, first_line='{first_text[:30]}'")

            formatter = DocumentFormatter(doc, config)
            formatter.run()

            output = io.BytesIO()
            doc.save(output)
            size_kb = output.tell() / 1024
            output.seek(0)

            original_name = file.filename
            if original_name.lower().endswith('.docx'):
                original_name = original_name[:-5]
            mode_label = config['name']

            # 用安全文件名（不含特殊字符）
            safe_name = re.sub(r'[\\/*?:"<>|]', '_', f"{original_name}_{mode_label}")
            output_name = f"{safe_name}.docx"

            print(f"[FORMAT] Done: {size_kb:.1f}KB, output: {output_name}")

            from flask import make_response
            response = make_response(output.read())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{output_name}"'
            response.headers['Content-Length'] = output.tell()
            return response
        except Exception as e:
            error_msg = str(e)
            print(f"[FORMAT] ERROR: {error_msg}")
            if 'zip' in error_msg.lower() or 'file format' in error_msg.lower():
                flash('文件格式错误：不是有效的 .docx 文件。'
                      '请用 Word 或 WPS 另存为 ".docx" 格式后再上传。', 'error')
            else:
                flash(f'处理出错: {error_msg}', 'error')
            return render_template('format.html')

    return render_template('format.html')
